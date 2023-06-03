import openpyxl
from docx import Document
from tkinter import Tk, filedialog, messagebox, Button

plantilla_path = ""
datos_path = ""
ruta_guardado = ""

# Función para cargar la plantilla de Word
def cargar_plantilla():
    global plantilla_path
    root = Tk()
    root.withdraw()
    plantilla_path = filedialog.askopenfilename(title="Seleccionar archivo de plantilla Word")
    root.destroy()

# Función para cargar los datos desde un archivo de Excel
def cargar_datos():
    global datos_path
    root = Tk()
    root.withdraw()
    datos_path = filedialog.askopenfilename(title="Seleccionar archivo de Excel con los datos")
    root.destroy()

# Función para seleccionar la ruta de guardado de los contratos generados
def seleccionar_ruta_guardado():
    global ruta_guardado
    root = Tk()
    root.withdraw()
    ruta_guardado = filedialog.askdirectory(title="Seleccionar la carpeta de guardado")
    root.destroy()

# Función para generar los contratos
def generar_contratos():
    try:
        # Cargar el archivo de Excel con los datos
        wb = openpyxl.load_workbook(datos_path)
        sheet = wb.active

        cambios = {}

        # Iterar sobre las filas del archivo de Excel y guardar los cambios requeridos para cada contrato
        for row in sheet.iter_rows(min_row=2, values_only=True):
            texto = row[0]
            reemplazo = row[1]
            id = row[2]
            nombre = row[3]

            if id not in cambios:
                cambios[id] = {}

            cambios[id][texto] = reemplazo
            cambios[id]['nombre'] = nombre

        # Generar los contratos para cada conjunto de cambios
        for id, cambio in cambios.items():
            contrato = Document(plantilla_path)

            # Reemplazar el texto en los párrafos del contrato
            for paragraph in contrato.paragraphs:
                for key, value in cambio.items():
                    if key != 'nombre' and key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, str(value))

            # Reemplazar el texto en las celdas de las tablas del contrato
            for table in contrato.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in cambio.items():
                            if key != 'nombre' and key in cell.text:
                                cell.text = cell.text.replace(key, str(value))

            # Guardar el contrato generado
            nombre_contrato = f"{cambio['nombre']}.docx"
            ruta_contrato = f"{ruta_guardado}/{nombre_contrato}"
            contrato.save(ruta_contrato)

        messagebox.showinfo("Generación de contratos", "Se generaron los contratos correctamente.")
    except Exception as e:
        messagebox.showerror("Error", f"Ocurrió un error al generar los contratos:\n{str(e)}")

def main():
    root = Tk()
    root.title("Generador de Contratos")

    # Botón para seleccionar la plantilla de Word
    btn_plantilla = Button(root, text="Seleccionar Plantilla", command=cargar_plantilla)
    btn_plantilla.pack()

    # Botón para seleccionar el archivo de Excel con los datos
    btn_datos = Button(root, text="Seleccionar Datos", command=cargar_datos)
    btn_datos.pack()

    # Botón para seleccionar la carpeta de guardado
    btn_ruta_guardado = Button(root, text="Seleccionar Carpeta de Guardado", command=seleccionar_ruta_guardado)
    btn_ruta_guardado.pack()

    # Botón para generar los contratos
    btn_generar = Button(root, text="Generar Contratos", command=generar_contratos)
    btn_generar.pack()

    root.geometry('500x400+380+100')

    root.mainloop()

if __name__ == "__main__":
    main()
